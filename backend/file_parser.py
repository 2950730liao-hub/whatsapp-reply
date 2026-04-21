"""
文件解析工具 - 从上传文件中提取文字内容
支持: PDF, Word (.docx), 纯文本 (.txt), 图片（存路径不提取文字）
"""
import os
import logging

logger = logging.getLogger(__name__)

# 图片扩展名
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp"}
# 文档扩展名
PDF_EXTENSIONS = {".pdf"}
WORD_EXTENSIONS = {".docx", ".doc"}
TEXT_EXTENSIONS = {".txt", ".md", ".csv"}


def extract_text_from_file(file_path: str, filename: str) -> str:
    """
    根据文件类型提取文字内容。
    - PDF: 使用 pdfplumber 提取
    - Word: 使用 python-docx 提取
    - 图片: 返回空字符串（AI 通过 file_url 引用）
    - 文本: 直接读取
    返回提取出的文字内容。
    """
    ext = os.path.splitext(filename)[1].lower()

    try:
        if ext in IMAGE_EXTENSIONS:
            # 图片不提取文字，返回占位描述
            return ""

        elif ext in PDF_EXTENSIONS:
            return _extract_pdf(file_path)

        elif ext in WORD_EXTENSIONS:
            return _extract_docx(file_path)

        elif ext in TEXT_EXTENSIONS:
            return _extract_text(file_path)

        else:
            logger.warning(f"不支持的文件类型: {ext}")
            return ""

    except Exception as e:
        logger.error(f"文件解析失败 [{filename}]: {e}")
        return ""


def get_doc_type(filename: str) -> str:
    """根据文件名返回文档类型标识"""
    ext = os.path.splitext(filename)[1].lower()
    if ext in IMAGE_EXTENSIONS:
        return "image"
    elif ext in PDF_EXTENSIONS:
        return "pdf"
    elif ext in WORD_EXTENSIONS:
        return "word"
    elif ext in TEXT_EXTENSIONS:
        return "text"
    return "file"


def _extract_pdf(file_path: str) -> str:
    """提取 PDF 文字"""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text.strip())
        return "\n\n".join(text_parts)
    except ImportError:
        logger.warning("pdfplumber 未安装，尝试使用 PyPDF2")
        try:
            import PyPDF2
            text_parts = []
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t.strip())
            return "\n\n".join(text_parts)
        except ImportError:
            return "[PDF 解析库未安装，请运行: pip install pdfplumber]"


def _extract_docx(file_path: str) -> str:
    """提取 Word 文字"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))
        return "\n".join(paragraphs)
    except ImportError:
        return "[python-docx 未安装，请运行: pip install python-docx]"


def _extract_text(file_path: str) -> str:
    """读取纯文本文件"""
    for encoding in ["utf-8", "gbk", "utf-16"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    return ""
